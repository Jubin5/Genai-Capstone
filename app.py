# app.py
"""
AI-Powered Contract & Policy Simplifier (Gemini Edition)
With enhanced Streamlit UI using HTML & CSS
"""

import os
import re
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import streamlit as st
import google.generativeai as genai
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# --------- CONFIGURE GEMINI ---------
genai.configure(api_key=os.getenv("GOOGLE_API_KEY", st.secrets.get("GOOGLE_API_KEY", "")))

# --------- PAGE CONFIG ---------
st.set_page_config(
    page_title="AI Contract & Policy Simplifier",
    layout="wide",
    page_icon="📜",
)

# --------- CUSTOM CSS ---------
st.markdown("""
<style>
/* Background & Fonts */
body, .stApp {
    background-color: #f4f7fb;
    font-family: 'Poppins', sans-serif;
}

/* Title Box */
.title-box {
    background: linear-gradient(90deg, #3b82f6, #6366f1);
    color: white;
    text-align: center;
    padding: 25px;
    border-radius: 12px;
    margin-bottom: 30px;
    box-shadow: 0px 4px 10px rgba(0,0,0,0.15);
}

/* Upload Section */
.upload-section {
    background-color: white;
    padding: 20px;
    border-radius: 10px;
    box-shadow: 0px 3px 8px rgba(0,0,0,0.1);
}

/* Summary Box */
textarea {
    font-size: 15px !important;
    font-family: 'Courier New', monospace;
}

/* Buttons */
.stDownloadButton button {
    background: linear-gradient(90deg, #2563eb, #1e40af);
    color: white !important;
    border: none;
    border-radius: 8px;
    padding: 10px 20px;
}
.stDownloadButton button:hover {
    background: linear-gradient(90deg, #1e3a8a, #1e40af);
}

/* Progress Bar */
div[data-testid="stProgress"] > div > div > div {
    background-color: #2563eb !important;
}

/* Search box */
input[type="text"] {
    border-radius: 8px;
    border: 1px solid #2563eb;
}

/* Headings */
h2, h3 {
    color: #1e3a8a;
    font-weight: 600;
}
</style>
""", unsafe_allow_html=True)

# --------- TITLE SECTION ---------
st.markdown("""
<div class="title-box">
    <h1>📜 AI-Powered Contract & Policy Simplifier</h1>
    <h4>Using Google Gemini for fast and reliable document simplification</h4>
    <p><i>Upload your contract or policy — Get an easy-to-read summary instantly.</i></p>
</div>
""", unsafe_allow_html=True)

# --------- FILE UPLOAD SECTION ---------
st.markdown('<div class="upload-section">', unsafe_allow_html=True)
uploaded_file = st.file_uploader("📂 Upload PDF or DOCX", type=["pdf", "docx"])
st.markdown('</div>', unsafe_allow_html=True)

# --------- PDF Extraction ---------
def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()

# --------- DOCX Extraction ---------
def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text.strip()

# --------- Cleaning ---------
def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    return text.strip()

# --------- Chunking ---------
def split_text(text: str, chunk_size: int = 1200, overlap: int = 200) -> list:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks

# --------- Simplification with Gemini ---------
def process_chunk_with_gemini(chunk: str, model_name: str = "gemini-1.5-flash") -> str:
    prompt = f"""
    You are a legal document simplifier. Read the following text and:
    1. Summarize it in plain, simple English.
    2. List obligations, rights, risks, penalties, and critical dates clearly.

    Text:
    {chunk}
    """
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    return response.text.strip()

# --------- Export Functions ---------
def export_to_pdf(input_text: str, output_path: str):
    styles = getSampleStyleSheet()
    story = []
    for line in input_text.split("\n"):
        if line.startswith("###"):
            style = styles["Heading2"]
            story.append(Paragraph(line.replace("###", "").strip(), style))
        else:
            story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 12))
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    doc.build(story)

def export_to_docx(input_text: str, output_path: str):
    doc = Document()
    for line in input_text.split("\n"):
        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=2)
        else:
            doc.add_paragraph(line)
    doc.save(output_path)

# --------- MAIN APP LOGIC ---------
if uploaded_file:
    input_path = Path(uploaded_file.name)
    with open(input_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if input_path.suffix.lower() == ".pdf":
        raw_text = extract_text_from_pdf(str(input_path))
    else:
        raw_text = extract_text_from_docx(str(input_path))

    cleaned_text = clean_text(raw_text)
    chunks = split_text(cleaned_text)

    st.success(f"✅ Extracted {len(chunks)} chunks from {input_path.name}")

    st.subheader("⚙️ Processing with Gemini Model...")
    summaries = []
    progress = st.progress(0)
    for i, chunk in enumerate(chunks, start=1):
        summary = process_chunk_with_gemini(chunk)
        summaries.append(f"### Summary of Chunk {i}\n{summary}\n")
        progress.progress(i / len(chunks))

    final_summary = "\n\n".join(summaries)

    # Display Results
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📄 Original Document Preview")
        st.text_area("Original Text", raw_text[:2000], height=400)

    with col2:
        st.subheader("✨ Simplified Summary")
        st.text_area("Simplified Output", final_summary[:5000], height=400)

        query = st.text_input("🔍 Search Summary (e.g., 'penalty')")
        if query:
            matches = [line for line in final_summary.split("\n") if query.lower() in line.lower()]
            st.write(f"Found {len(matches)} matches:")
            st.text("\n".join(matches[:20]))

    # Exports
    Path("Final_Summary_Report.txt").write_text(final_summary, encoding="utf-8")
    export_to_pdf(final_summary, "Final_Summary_Report.pdf")
    export_to_docx(final_summary, "Final_Summary_Report.docx")

    st.markdown("---")
    st.subheader("⬇️ Export Options")
    st.download_button("📄 Download TXT", final_summary, file_name="Final_Summary_Report.txt")
    with open("Final_Summary_Report.pdf", "rb") as f:
        st.download_button("📘 Download PDF", f, file_name="Final_Summary_Report.pdf")
    with open("Final_Summary_Report.docx", "rb") as f:
        st.download_button("📝 Download DOCX", f, file_name="Final_Summary_Report.docx")
