# export_report.py
"""
Step 4 (Part 2): Export merged summary report to PDF and Word
"""

from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# --------- Function to export as PDF ---------
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from xml.sax.saxutils import escape

def export_to_pdf(text, filename):
    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()
    story = []
    for line in text.split("\n"):
        safe_line = escape(line)  # Escapes <, >, &
        story.append(Paragraph(safe_line, styles["Normal"]))
    doc.build(story)



# --------- Function to export as DOCX ---------
def export_to_docx(input_text: str, output_path: str):
    doc = Document()
    for line in input_text.split("\n"):
        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=2)
        else:
            doc.add_paragraph(line)
    doc.save(output_path)


# --------- Main Program ---------
if __name__ == "__main__":
    input_file = Path("Final_Summary_Report.txt")

    if not input_file.exists():
        print("⚠️ Final_Summary_Report.txt not found. Run merge_summaries.py first.")
        exit()

    text = input_file.read_text(encoding="utf-8")

    # Export PDF
    export_to_pdf(text, "Final_Summary_Report.pdf")
    print("✅ Exported report as Final_Summary_Report.pdf")

    # Export DOCX
    export_to_docx(text, "Final_Summary_Report.docx")
    print("✅ Exported report as Final_Summary_Report.docx")
